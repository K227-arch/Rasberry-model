"""Generate the Dataset Format Guide as a Word document."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

doc = Document()

title = doc.add_heading("Dataset Format Guide", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Runyoro-English Translation Model Training Data Requirements")
doc.add_paragraph("")

# Section 1
doc.add_heading("1. Core Principle", level=1)
doc.add_paragraph(
    "Every row in the dataset must be one complete, natural sentence or phrase "
    "in Runyoro paired with its exact equivalent in English."
)
doc.add_paragraph(
    "The model learns grammar, tense, word order, and context FROM THE PATTERNS "
    "IN THE DATA ITSELF - not from labels or annotations."
)

# Section 2
doc.add_heading("2. What GOOD Training Data Looks Like", level=1)

doc.add_heading("2.1 Complete Sentence Pairs (Best for Grammar and Context)", level=2)
t = doc.add_table(rows=7, cols=2)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Runyoro"
t.rows[0].cells[1].text = "English"
good = [
    ("Abaana nibasome", "The children are studying"),
    ("Nyineeka afeera aka ye", "The house owner dies for his home"),
    ("Tukaija kare bulaijo", "We will come early tomorrow"),
    ("Omukazi akagura ebitooke omu katale", "The woman bought bananas at the market"),
    ("Enkuba neyija kutonnyera ijo", "It was going to rain yesterday"),
    ("Ekyaro kyaitu kikonyera abanyakuli", "Our village has always helped those in need"),
]
for i, (r, e) in enumerate(good, 1):
    t.rows[i].cells[0].text = r
    t.rows[i].cells[1].text = e
doc.add_paragraph(
    "Why: The model sees patterns across many sentences and learns grammar automatically."
)

doc.add_heading("2.2 Short Phrase Pairs (Good for Vocabulary in Context)", level=2)
t = doc.add_table(rows=6, cols=2)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Runyoro"
t.rows[0].cells[1].text = "English"
phrases = [
    ("ensi yaitu", "our country"),
    ("omusiri gwaitu", "our garden"),
    ("emiti mingi", "many trees"),
    ("omuhendo murungi", "a good price"),
    ("entebe ntukura", "a red chair"),
]
for i, (r, e) in enumerate(phrases, 1):
    t.rows[i].cells[0].text = r
    t.rows[i].cells[1].text = e

# Section 3
doc.add_heading("3. What BAD Training Data Looks Like", level=1)

doc.add_heading("3.1 Dictionary Definitions (NOT translations)", level=2)
t = doc.add_table(rows=4, cols=2)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Runyoro"
t.rows[0].cells[1].text = "English (BAD)"
bad = [
    ("fuuta", "To fill mouth with food"),
    ("jabuka", "To be chipped, broken"),
    ("jaagira", "To waddle, straddle, walk with legs wide open"),
]
for i, (r, e) in enumerate(bad, 1):
    t.rows[i].cells[0].text = r
    t.rows[i].cells[1].text = e
doc.add_paragraph(
    'Why bad: These are definitions, not translations. '
    'The model learns to output "To [verb]" for everything.'
)

doc.add_heading("3.2 Grammar Annotations Mixed In", level=2)
doc.add_paragraph(
    "NEVER include annotations like (v.i.), (v.t.), (n.), Adv., Oku- in the data. "
    "The model will try to reproduce them in output."
)

doc.add_heading("3.3 Multiple Translations in One Cell", level=2)
doc.add_paragraph(
    "NEVER put multiple unrelated sentences separated by semicolons in one cell. "
    "Each row = exactly one pair."
)

doc.add_heading("3.4 Single Words Without Context", level=2)
t = doc.add_table(rows=4, cols=2)
t.style = "Table Grid"
t.rows[0].cells[0].text = "BAD (no context)"
t.rows[0].cells[1].text = "GOOD (in context)"
t.rows[1].cells[0].text = "ekitabu = Book"
t.rows[1].cells[1].text = "Ekitabu kyange kiri ha meeza = My book is on the table"
t.rows[2].cells[0].text = "omuti = Tree"
t.rows[2].cells[1].text = "Omuti guno munene = This tree is big"
t.rows[3].cells[0].text = "amazzi = Water"
t.rows[3].cells[1].text = "Amazzi gaitu niganoga = Our water is clean"

# Section 4
doc.add_heading("4. How to Build Grammar Awareness", level=1)
doc.add_paragraph("The model learns grammar from seeing the same patterns repeated naturally.")

doc.add_heading("4.1 Tense Coverage - Same verb in multiple tenses", level=2)
t = doc.add_table(rows=6, cols=3)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Runyoro"
t.rows[0].cells[1].text = "English"
t.rows[0].cells[2].text = "Tense (reference only)"
tenses = [
    ("Ninkoora", "I am working", "Present Continuous"),
    ("Nkakoora", "I worked", "Past Simple"),
    ("Ninkaija kukoora", "I will work", "Future Simple"),
    ("Nakooire", "I have worked", "Present Perfect"),
    ("Nkaba ninkoora", "I was working", "Past Continuous"),
]
for i, (r, e, te) in enumerate(tenses, 1):
    t.rows[i].cells[0].text = r
    t.rows[i].cells[1].text = e
    t.rows[i].cells[2].text = te
doc.add_paragraph("Note: The Tense column is for YOUR reference. Do NOT include it in training data.")

doc.add_heading("4.2 Noun Class Agreement", level=2)
t = doc.add_table(rows=7, cols=2)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Runyoro"
t.rows[0].cells[1].text = "English"
nouns = [
    ("Omwana murungi", "The child is good"),
    ("Abaana barungi", "The children are good"),
    ("Ekitabu kirungi", "The book is good"),
    ("Ebitabu birungi", "The books are good"),
    ("Omuti murungi", "The tree is good"),
    ("Emiti mirungi", "The trees are good"),
]
for i, (r, e) in enumerate(nouns, 1):
    t.rows[i].cells[0].text = r
    t.rows[i].cells[1].text = e

doc.add_heading("4.3 Subject-Verb Agreement", level=2)
t = doc.add_table(rows=7, cols=2)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Runyoro"
t.rows[0].cells[1].text = "English"
subj = [
    ("Ninkoora", "I am working"),
    ("Nukoora", "You are working"),
    ("Nakoora", "He/she is working"),
    ("Tukoora", "We are working"),
    ("Mukoora", "You (plural) are working"),
    ("Bakoora", "They are working"),
]
for i, (r, e) in enumerate(subj, 1):
    t.rows[i].cells[0].text = r
    t.rows[i].cells[1].text = e

# Section 5
doc.add_heading("5. Dataset File Format", level=1)
doc.add_paragraph("Use CSV or XLSX format with two columns:")
doc.add_paragraph("Column A: Runyoro", style="List Bullet")
doc.add_paragraph("Column B: English", style="List Bullet")
doc.add_paragraph("")
doc.add_paragraph("Rules for each cell:")
for r in [
    "One sentence or phrase per cell - never multiple sentences with semicolons",
    "No annotations - no (v.i.), (n.), Adv., Oku-, etc.",
    "No numbering - no 1., 2), a) at the start",
    "Natural text only - write how a person would actually speak",
    "Consistent punctuation - periods at end of sentences",
    "No HTML or special formatting - plain text only",
]:
    doc.add_paragraph(r, style="List Bullet")

# Section 6
doc.add_heading("6. Minimum Dataset Size Guidelines", level=1)
t = doc.add_table(rows=6, cols=2)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Dataset Size"
t.rows[0].cells[1].text = "Expected Quality"
sizes = [
    ("< 5,000 pairs", "Poor - memorizes, does not generalize"),
    ("5,000 - 15,000", "Fair - handles common patterns"),
    ("15,000 - 50,000", "Good - generalizes to new sentences"),
    ("50,000 - 200,000", "Very good - handles nuance"),
    ("> 200,000 pairs", "Excellent - near-human quality"),
]
for i, (s, q) in enumerate(sizes, 1):
    t.rows[i].cells[0].text = s
    t.rows[i].cells[1].text = q
doc.add_paragraph("")
doc.add_paragraph("Current dataset: ~25,800 clean pairs - in the Good range.")

# Section 7
doc.add_heading("7. Quality Checklist Before Training", level=1)
for c in [
    "Each row has exactly one Runyoro text and one English text",
    "Both sides are complete (not fragments)",
    'No dictionary formatting ("To + verb" definitions)',
    "No grammar annotations or metadata",
    "No multiple translations crammed into one cell",
    "The Runyoro side is actually Runyoro (not Luganda, Rukiga, or Nyankole)",
    "The English side is natural English",
    "Sentences cover diverse tenses and structures",
    "Data covers multiple subject domains",
    "No excessive duplicates",
]:
    doc.add_paragraph(c, style="List Bullet")

# Section 8
doc.add_heading("8. Most Valuable Data to Add", level=1)
for v in [
    "Conversational sentences - everyday dialogue",
    "Proverbs with meaning-equivalent English",
    "News-style sentences - formal register",
    "Instructions/commands - imperative mood",
    "Questions - interrogative forms",
    "Complex sentences - relative clauses, conditionals",
    "Sentences with varied pronouns",
]:
    doc.add_paragraph(v, style="List Bullet")

# Save
out = Path(r"c:\Users\MarvinCliveTwesige\Desktop\PROJECTS\Rasberry-model\runyoro_nmt\docs\Dataset_Format_Guide.docx")
doc.save(str(out))
print(f"Done: {out}")
