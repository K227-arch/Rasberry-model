#!/usr/bin/env python3
"""
Generate a clear, easy-to-read Word document explaining exactly
how the data cleaning was done on the Runyoro-Rutooro / English files.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).parent.parent / "docs" / "Cleaning_Report.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

# ── colours ──────────────────────────────────────────────────────────────────
INDIGO = RGBColor(0x07, 0x02, 0x35)
TEAL   = RGBColor(0x00, 0x6a, 0x61)
GREY   = RGBColor(0x47, 0x46, 0x4f)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
RED    = RGBColor(0xba, 0x1a, 0x1a)
GREEN  = RGBColor(0x00, 0x6a, 0x61)

def bg(cell, hex_color):
    tc = cell._tc; p = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color); p.append(shd)

def h1(doc, text):
    p = doc.add_heading(text, 1)
    for r in p.runs: r.font.color.rgb = INDIGO
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)

def h2(doc, text):
    p = doc.add_heading(text, 2)
    for r in p.runs: r.font.color.rgb = TEAL
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)

def para(doc, text, bold=False, color=None, size=11, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent: p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold
    if color: r.font.color.rgb = color
    return p

def bullet_item(doc, label, detail, indent=0.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(indent)
    r1 = p.add_run(f"• {label}  ")
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = TEAL
    r2 = p.add_run(detail)
    r2.font.size = Pt(11); r2.font.color.rgb = GREY

def banner(doc, text, bg_hex="070235", fg=WHITE):
    t = doc.add_table(1, 1); t.style = "Table Grid"
    cell = t.rows[0].cells[0]; bg(cell, bg_hex)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = fg
    doc.add_paragraph()

def example_table(doc, rows_data):
    """rows_data: list of (runyoro, eng_before, eng_after)"""
    tbl = doc.add_table(rows=len(rows_data)+1, cols=3)
    tbl.style = "Table Grid"
    headers = ["Runyoro-Rutooro", "English BEFORE", "English AFTER"]
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        bg(cell, "006A61")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = WHITE; r.bold = True; r.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
    for i, (rny, before, after) in enumerate(rows_data, 1):
        row = tbl.rows[i]
        row.cells[0].text = rny
        row.cells[1].text = before
        row.cells[2].text = after
        if i % 2 == 0:
            bg(row.cells[0], "F2F4F6"); bg(row.cells[1], "F2F4F6")
        bg(row.cells[2], "EBF8F6")
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(10)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)
    doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(3.0); sec.right_margin = Cm(2.5)

# ── Title ────────────────────────────────────────────────────────────────────
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(20)
r = tp.add_run("Data Cleaning Report")
r.font.size = Pt(28); r.bold = True; r.font.color.rgb = INDIGO

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sp.add_run("Runyoro-Rutooro \u2194 English NMT \u2014 runyoro-nmt-v1")
r2.font.size = Pt(14); r2.font.color.rgb = TEAL

sp2 = doc.add_paragraph()
sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = sp2.add_run("How the raw data was cleaned before training the translation model")
r3.font.size = Pt(11); r3.font.color.rgb = GREY; r3.italic = True

doc.add_paragraph()
div = doc.add_table(1,1); bg(div.rows[0].cells[0], "006A61")
div.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Pt(2)
div.rows[0].cells[0].paragraphs[0].paragraph_format.space_after  = Pt(2)
doc.add_paragraph()
doc.add_page_break()

# ── Section 1: What is Cleaning ──────────────────────────────────────────────
h1(doc, "1.  What is Data Cleaning?")
para(doc,
    "Before an AI translation model can be trained, the raw text data needs to be "
    "prepared carefully. Cleaning means going through every sentence pair — one side "
    "in Runyoro-Rutooro and the other in English — and fixing small errors so the "
    "model learns from correct examples rather than mistakes."
)
para(doc,
    "Cleaning does NOT remove pairs. It fixes them in place. Every single pair from "
    "the source files stays in the dataset — only the text inside them is corrected "
    "where needed."
)
doc.add_paragraph()
banner(doc,
    "3,779 pairs were processed.  323 of them (8.5%) had something fixed.  "
    "The other 3,456 were already clean and passed through unchanged.",
    bg_hex="070235"
)

# ── Section 2: What was in the raw files ─────────────────────────────────────
h1(doc, "2.  What Was in the Raw Files?")
para(doc,
    "The data came from 8 source files — a mix of Excel spreadsheets and Word "
    "documents — containing Runyoro-Rutooro words, phrases, and sentences paired "
    "with their English translations."
)
doc.add_paragraph()
h2(doc, "The 8 source files")

src_tbl = doc.add_table(9, 3)
src_tbl.style = "Table Grid"
src_rows = [
    ("File", "Type", "Pairs extracted"),
    ("Agriculture Seed Vocabulary.csv.xlsx", "Spreadsheet", "123"),
    ("augmentted pos pairs.xlsx",            "Spreadsheet", "2,776"),
    ("Ff_fixed worked on.docx",              "Word document", "111"),
    ("J_fixed Worked on.docx",               "Word document", "165"),
    ("Tt fixed worked on.docx",              "Word document", "584"),
    ("U_fixed ...docx",                      "Word document", "1"),
    ("V_fixed worked on.docx",               "Word document", "4"),
    ("W_fixed worked on.docx",               "Word document", "18"),
]
for i, (a, b, c) in enumerate(src_rows):
    row = src_tbl.rows[i]
    row.cells[0].text = a; row.cells[1].text = b; row.cells[2].text = c
    if i == 0:
        for cell in row.cells:
            bg(cell, "070235")
            for p in cell.paragraphs:
                for r in p.runs: r.font.color.rgb = WHITE; r.bold = True; r.font.size = Pt(10)
    elif i % 2 == 0:
        for cell in row.cells: bg(cell, "F2F4F6")
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs: r.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
doc.add_paragraph()
para(doc, "Total: 3,782 raw pairs extracted across all 8 files.", bold=True, color=INDIGO)
doc.add_page_break()

# ── Section 3: The 8 cleaning steps ──────────────────────────────────────────
h1(doc, "3.  The 8 Cleaning Steps Applied")
para(doc,
    "Every pair was passed through 8 automatic cleaning steps in order. "
    "Each step checks for a specific type of problem and fixes it if found. "
    "A pair can be fixed by more than one step."
)
doc.add_paragraph()

steps = [
    (
        "Step 1: Unicode Normalisation",
        "233 pairs affected",
        "Some characters in the files were stored in slightly different formats — "
        "for example the same letter could be encoded in two different ways that "
        "look identical on screen but are different in the computer. "
        "This step converts all characters to a single standard form (NFC) so "
        "the computer treats identical-looking letters the same way.",
        "No visible change to the text — invisible technical fix only.",
        "006A61"
    ),
    (
        "Step 2: HTML Entity Decoding",
        "Rare — only where web-copied text was pasted",
        "Some text copied from websites contained HTML codes like &amp; (meaning &) "
        "or &quot; (meaning \"). This step converts them back to the real characters.",
        'Example:  &amp;  \u2192  &     |     &quot;  \u2192  "',
        "444173"
    ),
    (
        "Step 3: URL and Email Removal",
        "Very rare",
        "If any pair contained a web link or email address, it was removed. "
        "These do not help the model learn language and would confuse it.",
        "Example:  'Visit www.example.com for more'  \u2192  'Visit  for more'",
        "444173"
    ),
    (
        "Step 4: Quote Normalisation",
        "5 pairs (3 Runyoro side + 2 English side)",
        "Some text used curly or smart quotes (\u201clike this\u201d) instead of "
        "straight quotes (\"like this\"). This step converts all quote styles to "
        "straight quotes for consistency.",
        "\u201cOraire ota?\u201d  \u2192  \"Oraire ota?\"",
        "006A61"
    ),
    (
        "Step 5: Whitespace Normalisation",
        "1 pair",
        "Some entries had extra spaces, tabs, or line breaks inside the text. "
        "These were collapsed to a single space and leading/trailing spaces removed.",
        "Example:  'cassava   flour'  \u2192  'cassava flour'",
        "444173"
    ),
    (
        "Step 6: Leading Numbering Removed",
        "16 pairs (9 English + 7 Runyoro)",
        "Some entries were numbered in the original files, like '1. cassava' or "
        "'a) rice'. The number and punctuation at the start was removed because "
        "it is not part of the translation.",
        "Example:  '1. cassava flour'  \u2192  'cassava flour'\n"
        "          'a) okugesa'  \u2192  'okugesa'",
        "006A61"
    ),
    (
        "Step 7: Punctuation Spacing Fix",
        "90 pairs (88 English + 2 Runyoro)",
        "Some entries had a space before a punctuation mark — for example "
        "'harvesting .' or 'rice , beans'. The extra space was removed.",
        "Example:  'harvest time is a time of joy .'  \u2192  "
        "'harvest time is a time of joy.'",
        "444173"
    ),
    (
        "Step 8: English Capitalisation Fix",
        "233 pairs (the largest fix)",
        "This was the most common problem. English words and sentences in the data "
        "started with a lowercase letter when they should start with a capital. "
        "Every English entry that started with a lowercase letter was capitalised.\n\n"
        "This is a standard English rule: every sentence, word entry, and phrase "
        "should begin with a capital letter.",
        "Example:  'cassava'  \u2192  'Cassava'\n"
        "          'we grow cassava in our garden'  \u2192  'We grow cassava in our garden'",
        "006A61"
    ),
]

for i, (title, count, desc, example, color) in enumerate(steps):
    # Step header box
    st = doc.add_table(1, 1); st.style = "Table Grid"
    scell = st.rows[0].cells[0]; bg(scell, color)
    scell.paragraphs[0].clear()
    sp = scell.paragraphs[0]
    sp.paragraph_format.left_indent  = Cm(0.3)
    sp.paragraph_format.space_before = Pt(4)
    sp.paragraph_format.space_after  = Pt(4)
    r1 = sp.add_run(f"{title}     ")
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = WHITE
    r2 = sp.add_run(f"[{count}]")
    r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0xc4, 0xc1, 0xfb)

    # Description
    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.left_indent = Cm(0.5)
    p_desc.paragraph_format.space_after = Pt(2)
    rd = p_desc.add_run(desc)
    rd.font.size = Pt(11); rd.font.color.rgb = GREY

    # Example
    if example:
        p_ex = doc.add_paragraph()
        p_ex.paragraph_format.left_indent = Cm(0.8)
        p_ex.paragraph_format.space_after = Pt(8)
        re = p_ex.add_run(example)
        re.font.size = Pt(10); re.italic = True; re.font.color.rgb = TEAL

doc.add_page_break()

# ── Section 4: Summary table ──────────────────────────────────────────────────
h1(doc, "4.  Summary of All Fixes")
para(doc, "Here is a count of every type of fix applied across all 3,779 pairs:")
doc.add_paragraph()

sum_tbl = doc.add_table(10, 3)
sum_tbl.style = "Table Grid"
sum_rows = [
    ("Fix Type",                              "Which side", "Count"),
    ("Capitalisation fixed",                  "English",    "233"),
    ("Punctuation spacing fixed",             "English",    "88"),
    ("Leading number/letter removed",         "English",    "9"),
    ("Leading number/letter removed",         "Runyoro",    "7"),
    ("Quote style normalised",                "Runyoro",    "3"),
    ("Quote style normalised",                "English",    "2"),
    ("Repeated punctuation collapsed",        "Runyoro",    "2"),
    ("Extra whitespace removed",              "English",    "1"),
    ("TOTAL pairs with at least one fix",     "",           "323  (8.5% of 3,779)"),
]
for i, (a, b, c) in enumerate(sum_rows):
    row = sum_tbl.rows[i]
    row.cells[0].text = a; row.cells[1].text = b; row.cells[2].text = c
    if i == 0:
        for cell in row.cells:
            bg(cell, "070235")
            for p in cell.paragraphs:
                for r in p.runs: r.font.color.rgb = WHITE; r.bold = True; r.font.size = Pt(11)
    elif i == 9:
        for cell in row.cells:
            bg(cell, "006A61")
            for p in cell.paragraphs:
                for r in p.runs: r.font.color.rgb = WHITE; r.bold = True; r.font.size = Pt(11)
    elif i % 2 == 0:
        for cell in row.cells: bg(cell, "F2F4F6")
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs: r.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
doc.add_paragraph()

# Key point box
kp = doc.add_table(1, 1); kp.style = "Table Grid"
kpcell = kp.rows[0].cells[0]; bg(kpcell, "EBF8F6")
kpcell.paragraphs[0].clear()
kpp = kpcell.paragraphs[0]
kpp.paragraph_format.left_indent  = Cm(0.4)
kpp.paragraph_format.space_before = Pt(6); kpp.paragraph_format.space_after = Pt(6)
r_kp = kpp.add_run(
    "\u2139  Important note:  The Runyoro-Rutooro side was almost never changed.  "
    "Only 12 Runyoro entries were touched (numbering removal + quote fixes).  "
    "The Runyoro text was already well-formatted.  "
    "91.5% of fixes were on the English side."
)
r_kp.font.size = Pt(11); r_kp.font.color.rgb = INDIGO
doc.add_paragraph()
doc.add_page_break()

# ── Section 5: Before and after examples ─────────────────────────────────────
h1(doc, "5.  Real Before and After Examples")
para(doc, "These are actual examples from the data, showing what changed:")
doc.add_paragraph()

h2(doc, "5a.  Single Words — Capitalisation")
para(doc,
    "Even single words were kept and cleaned. The Runyoro word stays as-is; "
    "the English translation gets capitalised.",
    color=GREY
)
example_table(doc, [
    ("muhogo",     "cassava",    "Cassava"),
    ("omuceri",    "rice",       "Rice"),
    ("ekiyuni",    "yam",        "Yam"),
    ("ekikeke",    "pumpkin",    "Pumpkin"),
    ("kugesa",     "harvesting", "Harvesting"),
    ("ente",       "cows",       "Cows"),
    ("abantu",     "people",     "People"),
    ("amaizi",     "water",      "Water"),
])

h2(doc, "5b.  Short Phrases — Capitalisation")
example_table(doc, [
    ("ubuhunga bwa muhogo",           "cassava flour",                "Cassava flour"),
    ("ebikora bya muhogo",            "cassava leaves",               "Cassava leaves"),
    ("okugesa omubingo",              "harvesting millet",            "Harvesting millet"),
    ("omusiri gwa karo",              "today's garden",               "Today's garden"),
])

h2(doc, "5c.  Full Sentences — Capitalisation")
example_table(doc, [
    ("tulima muhogo mumusiri gwaitu",
     "we grow cassava in our garden",
     "We grow cassava in our garden"),
    ("tucumba omuceri buli kiro",
     "we cook rice every day",
     "We cook rice every day"),
    ("omuhendo gwo'muceri gweyongire",
     "the price of rice has increased",
     "The price of rice has increased"),
    ("amagesa gabaire marungi omwaka gunu",
     "the harvest was good this year",
     "The harvest was good this year"),
    ("nitugesa ebicoli kiro kinu",
     "we are harvesting maize today",
     "We are harvesting maize today"),
    ("hanyuma yo'kugesa twahura esingo",
     "after harvesting we store the grain",
     "After harvesting we store the grain"),
])
doc.add_page_break()

h2(doc, "5d.  Punctuation Spacing Fix")
para(doc,
    "Some entries had a space before the full stop or comma. The extra space was removed.",
    color=GREY
)
example_table(doc, [
    ("akasumi kokugesa",
     "harvest time is a time of joy .",
     "harvest time is a time of joy."),
    ("okugeza",
     "to test , to try",
     "to test, to try"),
    ("okukora",
     "to work ; to do",
     "to work; to do"),
])

h2(doc, "5e.  Leading Numbering Removed")
para(doc,
    "Entries that started with a number or letter from a list had the numbering stripped.",
    color=GREY
)
example_table(doc, [
    ("1. okugesa",   "1. harvesting",    "harvesting"),
    ("okugeza",      "a) to test",       "to test"),
    ("2. ente",      "2. cows",          "Cows"),
])

h2(doc, "5f.  Quote Normalisation")
para(doc, "Curly quotes were converted to straight quotes.", color=GREY)
example_table(doc, [
    ("\u201coraire ota?\u201d",  "\u201cHow are you?\u201d", '"How are you?"'),
    ("akagamba \u2018nkora\u2019", "he said \u2018I work\u2019", "he said 'I work'"),
])
doc.add_paragraph()
doc.add_page_break()

# ── Section 6: What was NOT changed ──────────────────────────────────────────
h1(doc, "6.  What Was NOT Changed")
para(doc,
    "It is just as important to understand what the cleaning did NOT do, "
    "to make sure the data stays faithful to the original:"
)
doc.add_paragraph()

not_changed = [
    ("Runyoro-Rutooro words or spellings",
     "The Runyoro side was not corrected, edited, or changed in any way "
     "except removing accidental leading numbers. The language, spelling, "
     "and word choice were left exactly as the human authors wrote them."),
    ("Short or single-word pairs",
     "Single words like 'muhogo' (cassava) or 'ente' (cows) were kept. "
     "These are valid translation pairs and help the model learn vocabulary."),
    ("Long pairs",
     "Long sentences were also kept. There was no maximum length filter."),
    ("Number and symbol content",
     "Pairs that included numbers or symbols alongside text were kept."),
    ("Pairs with unusual length ratios",
     "Even if the Runyoro side is much shorter or longer than the English side, "
     "the pair was kept. This is natural in translation — Runyoro is agglutinative "
     "and one word can mean a whole English phrase."),
    ("Meaning or word choice",
     "The cleaning never changed the meaning of any translation. It only fixed "
     "formatting — capitalisation, spacing, quotes, and numbering."),
]

for title, detail in not_changed:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(f"\u2022  {title}\n")
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = INDIGO
    r2 = p.add_run(f"     {detail}")
    r2.font.size = Pt(11); r2.font.color.rgb = GREY

doc.add_paragraph()
doc.add_page_break()

# ── Section 7: Final data count ───────────────────────────────────────────────
h1(doc, "7.  How Much Data Is Left After Cleaning")
para(doc,
    "Cleaning does not remove data — it only improves it. "
    "Here is the full picture of what remained after cleaning:"
)
doc.add_paragraph()

count_tbl = doc.add_table(7, 3)
count_tbl.style = "Table Grid"
count_rows = [
    ("Stage",                          "Count",  "Notes"),
    ("Raw pairs from 8 files",         "3,782",  "Before any processing"),
    ("After removing duplicates only", "3,779",  "Only 3 exact duplicates removed"),
    ("After cleaning",                 "3,779",  "323 fixed, 0 removed"),
    ("After augmentation (+923 new)",  "4,702",  "New pairs created from originals"),
    ("Training set (85%)",             "3,996",  "Used to teach the model"),
    ("Validation + Test (15%)",        "470 + 236", "Used to measure quality"),
]
for i, (a, b, c) in enumerate(count_rows):
    row = count_tbl.rows[i]
    row.cells[0].text = a; row.cells[1].text = b; row.cells[2].text = c
    if i == 0:
        for cell in row.cells:
            bg(cell, "070235")
            for p in cell.paragraphs:
                for r in p.runs: r.font.color.rgb = WHITE; r.bold = True; r.font.size = Pt(11)
    elif i == 3:
        for cell in row.cells:
            bg(cell, "EBF8F6")
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(11)
    elif i == 4:
        for cell in row.cells:
            bg(cell, "006A61")
            for p in cell.paragraphs:
                for r in p.runs: r.font.color.rgb = WHITE; r.bold = True; r.font.size = Pt(11)
    elif i % 2 == 0:
        for cell in row.cells: bg(cell, "F2F4F6")
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs: r.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)

doc.add_paragraph()

# Final summary box
fs = doc.add_table(1, 1); fs.style = "Table Grid"
fscell = fs.rows[0].cells[0]; bg(fscell, "070235")
fscell.paragraphs[0].clear()
lines = [
    ("Final Summary", True, 14, WHITE),
    ("", False, 5, WHITE),
    ("Total pairs from raw files:          3,782", False, 11, RGBColor(0x86,0xf2,0xe4)),
    ("Pairs removed (duplicates only):         3", False, 11, RGBColor(0xc4,0xc1,0xfb)),
    ("Pairs cleaned and fixed:               323  (8.5%)", False, 11, RGBColor(0x86,0xf2,0xe4)),
    ("Pairs unchanged (already clean):     3,456  (91.5%)", False, 11, RGBColor(0xc4,0xc1,0xfb)),
    ("", False, 5, WHITE),
    ("TOTAL clean pairs ready:             3,779", True, 13, WHITE),
    ("  \u2514\u2500 Training:    3,996 pairs  (with augmentation)", False, 11, RGBColor(0x86,0xf2,0xe4)),
    ("  \u2514\u2500 Validation:   470 pairs", False, 11, RGBColor(0x86,0xf2,0xe4)),
    ("  \u2514\u2500 Test:          236 pairs", False, 11, RGBColor(0x86,0xf2,0xe4)),
    ("", False, 5, WHITE),
    ("Model trained:  NLLB-200 1.3B   |   BLEU = 18.77   |   chrF++ = 22.53", True, 11, WHITE),
]
for j, (text, bold, size, color) in enumerate(lines):
    p = fscell.paragraphs[0] if j == 0 else fscell.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    if text:
        r = p.add_run(text)
        r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color

p_end = fscell.add_paragraph(); p_end.paragraph_format.space_after = Pt(6)
doc.add_paragraph()

# Footer
pf = doc.add_paragraph()
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = pf.add_run(
    "Generated by runyoro-nmt-v1 pipeline  \u2022  "
    "kathay/runyoro-nmt-v1  \u2022  "
    "huggingface.co/spaces/kathay/runyoro-translator"
)
rf.font.size = Pt(9); rf.font.color.rgb = GREY; rf.italic = True

# ── Save ─────────────────────────────────────────────────────────────────────
doc.save(str(OUT))
print(f"Saved: {OUT}")
