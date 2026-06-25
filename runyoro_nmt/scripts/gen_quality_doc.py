"""Generate Dataset Quality Assessment document."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

doc = Document()
title = doc.add_heading("Dataset Quality Assessment & Improvement Guide", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Runyoro-English NMT Model v2 Training Data Analysis")
doc.add_paragraph("")

doc.add_heading("1. Overall Rating: 6/10", level=1)
doc.add_paragraph(
    "The current dataset (25,832 pairs) is in the Good range for a low-resource "
    "language, but has clear gaps that limit translation quality."
)

doc.add_heading("2. Strengths", level=1)
for s in [
    "25,832 pairs is in the Good range (15K-50K threshold)",
    "Sentence data from Supabase is high quality - proper human translations",
    "Covers multiple domains (agriculture, religion, daily life, history, nature)",
    "Bidirectional training doubles effective samples to 51,664",
    "Data was cleaned: no dictionary definitions, no grammar annotations",
]:
    doc.add_paragraph(s, style="List Bullet")

doc.add_heading("3. Weaknesses", level=1)

doc.add_heading("3.1 Too Many Word Pairs, Not Enough Sentences", level=2)
doc.add_paragraph(
    "About 56% of the data are word/phrase-level pairs. Only 44% (11,442) are full "
    "sentences. The model needs more complete sentences to learn grammar and word order."
)
doc.add_paragraph("Example of what we have too much of:")
t = doc.add_table(rows=4, cols=2)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Runyoro"
t.rows[0].cells[1].text = "English"
t.rows[1].cells[0].text = "omuceri"
t.rows[1].cells[1].text = "Rice"
t.rows[2].cells[0].text = "empunu"
t.rows[2].cells[1].text = "Pig"
t.rows[3].cells[0].text = "okugura"
t.rows[3].cells[1].text = "Buy"
doc.add_paragraph("")
doc.add_paragraph("What we need MORE of:")
t = doc.add_table(rows=4, cols=2)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Runyoro"
t.rows[0].cells[1].text = "English"
t.rows[1].cells[0].text = "Tugura omuceri omu katale buli wiiki"
t.rows[1].cells[1].text = "We buy rice at the market every week"
t.rows[2].cells[0].text = "Empunu zaitu zikaire"
t.rows[2].cells[1].text = "Our pigs are healthy"
t.rows[3].cells[0].text = "Nkagura ebitooke ijo"
t.rows[3].cells[1].text = "I bought bananas yesterday"

doc.add_heading("3.2 English-First Bias", level=2)
doc.add_paragraph(
    "Most pairs were created by writing English first, then translating to Runyoro. "
    "This makes the model better at English-to-Runyoro than Runyoro-to-English."
)
doc.add_paragraph("What we need: Sentences originally written in Runyoro:")
t = doc.add_table(rows=5, cols=2)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Runyoro (original)"
t.rows[0].cells[1].text = "English (translation)"
t.rows[1].cells[0].text = "Nyineeka afeera aka ye"
t.rows[1].cells[1].text = "The house owner dies for his home"
t.rows[2].cells[0].text = "Akaguru kakukora gye omuhanda"
t.rows[2].cells[1].text = "The leg took him/her along the road"
t.rows[3].cells[0].text = "Obutamu bwensoni tiburungi"
t.rows[3].cells[1].text = "Shamelessness is not good"
t.rows[4].cells[0].text = "Enjoka etarikuuma tekubwa"
t.rows[4].cells[1].text = "A snake that does not hiss does not bite"

doc.add_heading("3.3 Limited Tense Coverage", level=2)
doc.add_paragraph(
    "The data lacks diversity in verb tenses. The model struggles with past and "
    "future because it has not seen enough examples."
)
doc.add_paragraph("What we need - same concept in multiple tenses:")
t = doc.add_table(rows=6, cols=3)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Runyoro"
t.rows[0].cells[1].text = "English"
t.rows[0].cells[2].text = "Tense"
t.rows[1].cells[0].text = "Ninkoora"
t.rows[1].cells[1].text = "I am working"
t.rows[1].cells[2].text = "Present"
t.rows[2].cells[0].text = "Nkakoora"
t.rows[2].cells[1].text = "I worked"
t.rows[2].cells[2].text = "Past"
t.rows[3].cells[0].text = "Ninkaija kukoora"
t.rows[3].cells[1].text = "I will work"
t.rows[3].cells[2].text = "Future"
t.rows[4].cells[0].text = "Nakooire"
t.rows[4].cells[1].text = "I have worked"
t.rows[4].cells[2].text = "Perfect"
t.rows[5].cells[0].text = "Nkaba ninkoora"
t.rows[5].cells[1].text = "I was working"
t.rows[5].cells[2].text = "Past Continuous"

doc.add_heading("3.4 Missing Conversational Data", level=2)
doc.add_paragraph(
    "The dataset has almost no everyday dialogue. The model cannot handle "
    "greetings, questions, or short exchanges."
)
doc.add_paragraph("Examples of conversational pairs needed:")
t = doc.add_table(rows=9, cols=2)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Runyoro"
t.rows[0].cells[1].text = "English"
t.rows[1].cells[0].text = "Osiibire ota?"
t.rows[1].cells[1].text = "How are you?"
t.rows[2].cells[0].text = "Nsiibire kurungi, naiwe?"
t.rows[2].cells[1].text = "I am fine, and you?"
t.rows[3].cells[0].text = "Orikugenda he?"
t.rows[3].cells[1].text = "Where are you going?"
t.rows[4].cells[0].text = "Ninkugenda omu katale"
t.rows[4].cells[1].text = "I am going to the market"
t.rows[5].cells[0].text = "Webale muno"
t.rows[5].cells[1].text = "Thank you very much"
t.rows[6].cells[0].text = "Tukaizana"
t.rows[6].cells[1].text = "See you later"
t.rows[7].cells[0].text = "Ninkusaba onyambe"
t.rows[7].cells[1].text = "Please help me"
t.rows[8].cells[0].text = "Oizire ryari?"
t.rows[8].cells[1].text = "When did you come?"

doc.add_heading("3.5 Domain Gaps", level=2)
doc.add_paragraph("Some important domains have very little coverage:")
t = doc.add_table(rows=6, cols=3)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Domain"
t.rows[0].cells[1].text = "Status"
t.rows[0].cells[2].text = "Example Pair Needed"
t.rows[1].cells[0].text = "Technology"
t.rows[1].cells[1].text = "Almost none"
t.rows[1].cells[2].text = "Simu yange eyahirize batiri = My phone battery is dead"
t.rows[2].cells[0].text = "Health/Medical"
t.rows[2].cells[1].text = "Very few"
t.rows[2].cells[2].text = "Omwana araine obushwere = The child has malaria"
t.rows[3].cells[0].text = "Commerce"
t.rows[3].cells[1].text = "Limited"
t.rows[3].cells[2].text = "Omuhendo nigugye haiguru = The price has gone up"
t.rows[4].cells[0].text = "Education"
t.rows[4].cells[1].text = "Some"
t.rows[4].cells[2].text = "Abaisiki nibasoma kurungi = The girls study well"
t.rows[5].cells[0].text = "Family"
t.rows[5].cells[1].text = "Few"
t.rows[5].cells[2].text = "Maama wange akaija ijo = My mother came yesterday"

doc.add_heading("4. Roadmap to 8/10", level=1)
t = doc.add_table(rows=6, cols=3)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Action"
t.rows[0].cells[1].text = "Pairs to Add"
t.rows[0].cells[2].text = "Priority"
t.rows[1].cells[0].text = "Add 10K+ sentence pairs (Supabase contributors)"
t.rows[1].cells[1].text = "10,000"
t.rows[1].cells[2].text = "HIGH"
t.rows[2].cells[0].text = "Add 2K conversational pairs"
t.rows[2].cells[1].text = "2,000"
t.rows[2].cells[2].text = "HIGH"
t.rows[3].cells[0].text = "Add Runyoro proverbs with translations"
t.rows[3].cells[1].text = "500"
t.rows[3].cells[2].text = "MEDIUM"
t.rows[4].cells[0].text = "Add tense-varied sentences"
t.rows[4].cells[1].text = "3,000"
t.rows[4].cells[2].text = "HIGH"
t.rows[5].cells[0].text = "Fill domain gaps (tech, health, commerce)"
t.rows[5].cells[1].text = "5,000"
t.rows[5].cells[2].text = "MEDIUM"
doc.add_paragraph("")
doc.add_paragraph(
    "Target: 50,000+ clean sentence pairs for production-quality translation."
)

out = Path(
    r"c:\Users\MarvinCliveTwesige\Desktop\PROJECTS\Rasberry-model"
    r"\runyoro_nmt\docs\Dataset_Quality_Assessment.docx"
)
doc.save(str(out))
print(f"Done: {out}")
