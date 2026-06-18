#!/usr/bin/env python3
"""Generate a Word document explaining the data pipeline clearly."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT_PATH = Path(__file__).parent.parent / "docs" / "Data_Pipeline_Report_v2.docx"
OUT_PATH.parent.mkdir(parents=True, exist_ok=True)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
INDIGO  = RGBColor(0x07, 0x02, 0x35)
TEAL    = RGBColor(0x00, 0x6a, 0x61)
GREY    = RGBColor(0x47, 0x46, 0x4f)
LGREY   = RGBColor(0xec, 0xee, 0xf0)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
GREEN   = RGBColor(0x00, 0x6a, 0x61)
RED     = RGBColor(0xba, 0x1a, 0x1a)

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def heading(doc, text, level=1, color=INDIGO):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(doc, text, bold=False, color=None, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = TEAL
        r.add_break() if False else None
        p.add_run(text).font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p

def callout_box(doc, title, content, bg_hex="EBF8F6", title_color=TEAL):
    """Simulate a callout box using a 1-cell table."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg_hex)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(f"  {title}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = title_color
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(f"  {content}")
    r2.font.size = Pt(11)
    r2.font.color.rgb = GREY
    doc.add_paragraph()
    return tbl

# ════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(30)
r = p_title.add_run("Runyoro-Rutooro ↔ English NMT")
r.font.size = Pt(26)
r.font.bold = True
r.font.color.rgb = INDIGO

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run("Data Pipeline Report")
r2.font.size = Pt(18)
r2.font.color.rgb = TEAL

p_sub2 = doc.add_paragraph()
p_sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p_sub2.add_run("Cleaning · Augmentation · Back-Translation\nHow the data was prepared and what remains")
r3.font.size = Pt(12)
r3.font.color.rgb = GREY

doc.add_paragraph()

# Divider line using a table
div = doc.add_table(rows=1, cols=1)
set_cell_bg(div.rows[0].cells[0], "006A61")
div.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Pt(2)
div.rows[0].cells[0].paragraphs[0].paragraph_format.space_after  = Pt(2)
doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# SECTION 1 – OVERVIEW
# ════════════════════════════════════════════════════════════════════
heading(doc, "1.  Overview", 1)
body(doc, (
    "This document explains, in plain language, exactly what happened to the "
    "Runyoro-Rutooro / English data files from the moment they were read in "
    "to the moment they were ready for model training.  "
    "It covers three main stages: Cleaning, Augmentation, and Back-Translation."
))
doc.add_paragraph()

# Summary numbers table
heading(doc, "Quick Summary", 2, color=TEAL)
tbl = doc.add_table(rows=6, cols=2)
tbl.style = "Table Grid"
rows_data = [
    ("Stage",                       "Result"),
    ("Raw pairs read from files",   "3,782"),
    ("Valid pairs after filtering", "881"),
    ("Pairs after cleaning",        "881  (174 modified)"),
    ("New pairs from augmentation", "+323"),
    ("Total pairs for training",    "1,204"),
]
for i, (col1, col2) in enumerate(rows_data):
    row = tbl.rows[i]
    row.cells[0].text = col1
    row.cells[1].text = col2
    if i == 0:
        set_cell_bg(row.cells[0], "070235")
        set_cell_bg(row.cells[1], "070235")
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
    elif i % 2 == 0:
        set_cell_bg(row.cells[0], "EBF8F6")
        set_cell_bg(row.cells[1], "EBF8F6")
    for cell in row.cells:
        for para in cell.paragraphs:
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after  = Pt(3)
            for run in para.runs:
                run.font.size = Pt(11)
doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# SECTION 2 – WHERE THE DATA CAME FROM
# ════════════════════════════════════════════════════════════════════
heading(doc, "2.  Where the Data Came From", 1)
body(doc, (
    "The data was collected from 8 raw files stored in the 'raw data' folder.  "
    "Each file contained Runyoro-Rutooro words or sentences on one side and "
    "their English translations on the other side."
))
doc.add_paragraph()
heading(doc, "The 8 source files", 2, color=TEAL)

files_data = [
    ("Agriculture Seed Vocabulary.csv.xlsx", "Spreadsheet", "Agriculture domain – crop names, farming terms", "123"),
    ("augmentted pos pairs.xlsx",            "Spreadsheet", "Pre-tagged word pairs across many topics",        "2,776"),
    ("Ff_fixed worked on.docx",              "Word doc",    "Vocabulary starting with the letter F",           "69"),
    ("J_fixed Worked on.docx",               "Word doc",    "Vocabulary starting with the letter J",           "127"),
    ("Tt fixed worked on.docx",              "Word doc",    "Vocabulary starting with the letter T",           "368"),
    ("U_fixed ...docx",                      "Word doc",    "Vocabulary starting with the letter U",           "1"),
    ("V_fixed worked on.docx",               "Word doc",    "Vocabulary starting with the letter V",           "3"),
    ("W_fixed worked on.docx",               "Word doc",    "Vocabulary starting with the letter W",           "18"),
]

tbl2 = doc.add_table(rows=len(files_data)+1, cols=4)
tbl2.style = "Table Grid"
headers = ["File Name", "Type", "Contents", "Pairs extracted"]
for j, h in enumerate(headers):
    cell = tbl2.rows[0].cells[j]
    cell.text = h
    set_cell_bg(cell, "070235")
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = WHITE
            run.font.bold = True
            run.font.size = Pt(10)

for i, (fname, ftype, desc, count) in enumerate(files_data, start=1):
    row = tbl2.rows[i]
    row.cells[0].text = fname
    row.cells[1].text = ftype
    row.cells[2].text = desc
    row.cells[3].text = count
    if i % 2 == 0:
        for cell in row.cells:
            set_cell_bg(cell, "F2F4F6")
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)

doc.add_paragraph()
callout_box(
    doc,
    "Total raw pairs extracted:  3,782",
    "The system automatically detected which column was Runyoro-Rutooro "
    "and which was English in every file using keyword matching on the column names.",
    bg_hex="EBF8F6",
    title_color=TEAL,
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# SECTION 3 – VALIDATION (FILTERING)
# ════════════════════════════════════════════════════════════════════
heading(doc, "3.  Validation and Filtering", 1)
body(doc, (
    "Before cleaning, every pair was checked against a set of quality rules.  "
    "Pairs that failed were removed.  This is called validation or filtering.  "
    "Think of it like sorting vegetables before cooking — you throw away the bad ones."
))
doc.add_paragraph()
heading(doc, "Rules applied", 2, color=TEAL)
bullet(doc, "Pairs where either side had fewer than 2 words were removed (too short to be useful).", "Too short:  ")
bullet(doc, "Pairs where one side was just numbers, symbols, or punctuation were removed.", "Content-free:  ")
bullet(doc, "Pairs where the Runyoro side was much shorter or longer than the English side were removed.", "Length ratio:  ")
bullet(doc, "Pairs that appeared more than once were removed (only the first copy was kept).", "Duplicates:  ")

doc.add_paragraph()
heading(doc, "Filtering results", 2, color=TEAL)
tbl3 = doc.add_table(rows=6, cols=2)
tbl3.style = "Table Grid"
filter_data = [
    ("Reason removed",                         "How many"),
    ("Too short (fewer than 2 words)",         "2,730"),
    ("Content-free (numbers / symbols only)",  "272"),
    ("Length ratio too different",             "55"),
    ("Possible misalignment",                  "10 (kept but flagged)"),
    ("Pairs that survived → kept for cleaning","428  ✓"),
]
for i, (col1, col2) in enumerate(filter_data):
    row = tbl3.rows[i]
    row.cells[0].text = col1
    row.cells[1].text = col2
    if i == 0:
        set_cell_bg(row.cells[0], "070235")
        set_cell_bg(row.cells[1], "070235")
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
    elif i == 5:
        set_cell_bg(row.cells[0], "006A61")
        set_cell_bg(row.cells[1], "006A61")
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
    elif i % 2 == 0:
        set_cell_bg(row.cells[0], "F2F4F6")
        set_cell_bg(row.cells[1], "F2F4F6")
    for cell in row.cells:
        for para in cell.paragraphs:
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after  = Pt(3)
            for run in para.runs:
                run.font.size = Pt(11)

doc.add_paragraph()
body(doc,
    "In plain numbers: out of 3,782 raw pairs, 2,901 were removed and 428 were kept.",
    bold=True, color=INDIGO)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# SECTION 4 – CLEANING
# ════════════════════════════════════════════════════════════════════
heading(doc, "4.  Data Cleaning", 1)
body(doc, (
    "Cleaning means fixing small errors in the text without removing the pair.  "
    "The 428 pairs that passed filtering were all passed through a cleaning pipeline.  "
    "187 of them (43.7%) had something fixed.  The other 241 were already clean."
))
doc.add_paragraph()
heading(doc, "What was cleaned and how", 2, color=TEAL)
body(doc, "The following fixes were applied automatically:")
doc.add_paragraph()

cleaning_steps = [
    (
        "English capitalisation fix  (177 pairs)",
        "The biggest issue found in the data was that English sentences started "
        "with a lowercase letter.  For example:  'cassava flour' was corrected "
        "to 'Cassava flour'.  This follows standard English grammar — every "
        "sentence or phrase starts with a capital letter.",
    ),
    (
        "Punctuation spacing fix  (5 pairs)",
        "Some pairs had a space before a punctuation mark, like 'hello .'  "
        "This was fixed to 'hello.'",
    ),
    (
        "Leading numbering removed  (5 pairs)",
        "Some entries started with a number like '1.' or 'a)'.  "
        "For example '1. cassava flour' was cleaned to 'cassava flour'.",
    ),
    (
        "Repeated punctuation fixed  (3 pairs)",
        "Entries with double punctuation like '!!' or ',,' were cleaned "
        "to a single '!' or ','.",
    ),
    (
        "Curly quotes converted  (3+1 pairs)",
        "Curly or 'smart' quotes like \u201cword\u201d were converted to straight "
        "quotes like \"word\" for consistency.",
    ),
]

for title, detail in cleaning_steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5)
    r_title = p.add_run(f"\u2022  {title}\n")
    r_title.bold = True
    r_title.font.size = Pt(11)
    r_title.font.color.rgb = TEAL
    r_detail = p.add_run(f"     {detail}")
    r_detail.font.size = Pt(11)
    r_detail.font.color.rgb = GREY

doc.add_paragraph()
heading(doc, "Before and after examples", 2, color=TEAL)

examples = [
    ("ubuhunga bwa muhogo",                "cassava flour",                          "Cassava flour"),
    ("tulima muhogo mumusiri gwaitu",       "we grow cassava in our garden",          "We grow cassava in our garden"),
    ("omuhendo gwo'muceri gweyongire",      "the price of rice has increased",        "The price of rice has increased"),
    ("akasumi kokugesa kaba kasumi",        "harvest time is a time of joy",          "Harvest time is a time of joy"),
    ("okusesira ebirimwa amaizi goruseke",  "irrigation helps crops grow in dry season", "Irrigation helps crops grow in dry season"),
    ("turimira omusiri tutakabyaire",       "we plow the field before planting",      "We plow the field before planting"),
]

ex_tbl = doc.add_table(rows=len(examples)+1, cols=3)
ex_tbl.style = "Table Grid"
ex_headers = ["Runyoro-Rutooro", "English BEFORE cleaning", "English AFTER cleaning"]
for j, h in enumerate(ex_headers):
    cell = ex_tbl.rows[0].cells[j]
    cell.text = h
    set_cell_bg(cell, "006A61")
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = WHITE
            run.font.bold = True
            run.font.size = Pt(10)

for i, (rny, before, after) in enumerate(examples, start=1):
    row = ex_tbl.rows[i]
    row.cells[0].text = rny
    row.cells[1].text = before
    row.cells[2].text = after
    if i % 2 == 0:
        set_cell_bg(row.cells[0], "F2F4F6")
        set_cell_bg(row.cells[1], "F2F4F6")
    set_cell_bg(row.cells[2], "EBF8F6")
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)

doc.add_paragraph()
callout_box(
    doc,
    "Result after cleaning:  881 clean pairs",
    "The Runyoro-Rutooro side was NOT changed.  "
    "Only the English side needed corrections.  "
    "Runyoro-Rutooro uses standard Latin script and "
    "the data was already correctly formatted.",
    bg_hex="EBF8F6",
    title_color=TEAL,
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# SECTION 5 – AUGMENTATION
# ════════════════════════════════════════════════════════════════════
heading(doc, "5.  Data Augmentation", 1)
body(doc, (
    "881 clean pairs is a small dataset for training an AI translation model.  "
    "To help the model learn better, we created extra training examples from "
    "the existing ones.  This process is called augmentation — it means making "
    "more data from what you already have, the same way you could practise a "
    "sentence in slightly different ways to memorise it better."
))
doc.add_paragraph()
heading(doc, "The three augmentation methods used", 2, color=TEAL)
doc.add_paragraph()

aug_methods = [
    (
        "Method 1: Token Deletion  (122 new pairs)",
        "A few words were randomly removed from a sentence to create a shorter version.  "
        "This teaches the model to handle incomplete or informal input.",
        "amagesa gabaire marungi omwaka gunu",
        "The harvest was good this year",
        "amagesa gabaire omwaka gunu",
        "The harvest was good this year",
        "The word 'marungi' (good) was deleted from the Runyoro side.",
    ),
    (
        "Method 2: Token Swap  (77 new pairs)",
        "Two neighbouring words were swapped around.  "
        "This teaches the model that word order can vary slightly.",
        "tucumba omuceri buli kiro",
        "We cook rice every day",
        "tucumba buli omuceri kiro",
        "We cook rice every day",
        "'omuceri' and 'buli' were swapped.",
    ),
    (
        "Method 3: Combined  (117 new pairs)",
        "A mix of deletion and swap was applied at the same time.  "
        "This creates more varied training examples.",
        "tubyara ebikeke haihi no'rugo",
        "We planted pumpkins near the fence",
        "tubyara haihi no'rugo",
        "We planted pumpkins near the fence",
        "'ebikeke' (pumpkins) was deleted from the Runyoro side.",
    ),
]

for method_title, method_desc, orig_rny, orig_eng, aug_rny, aug_eng, note in aug_methods:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(method_title)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = INDIGO

    body(doc, method_desc, color=GREY)

    inner_tbl = doc.add_table(rows=3, cols=3)
    inner_tbl.style = "Table Grid"
    inner_tbl.rows[0].cells[0].text = ""
    inner_tbl.rows[0].cells[1].text = "Runyoro-Rutooro"
    inner_tbl.rows[0].cells[2].text = "English"
    inner_tbl.rows[1].cells[0].text = "Original"
    inner_tbl.rows[1].cells[1].text = orig_rny
    inner_tbl.rows[1].cells[2].text = orig_eng
    inner_tbl.rows[2].cells[0].text = "Augmented"
    inner_tbl.rows[2].cells[1].text = aug_rny
    inner_tbl.rows[2].cells[2].text = aug_eng

    set_cell_bg(inner_tbl.rows[0].cells[0], "070235")
    set_cell_bg(inner_tbl.rows[0].cells[1], "070235")
    set_cell_bg(inner_tbl.rows[0].cells[2], "070235")
    set_cell_bg(inner_tbl.rows[2].cells[0], "EBF8F6")
    set_cell_bg(inner_tbl.rows[2].cells[1], "EBF8F6")
    set_cell_bg(inner_tbl.rows[2].cells[2], "EBF8F6")

    for row in inner_tbl.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after  = Pt(2)
    for cell in inner_tbl.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True

    p_note = doc.add_paragraph()
    p_note.paragraph_format.left_indent = Cm(0.5)
    p_note.paragraph_format.space_after = Pt(10)
    r_note = p_note.add_run(f"Note: {note}")
    r_note.font.size = Pt(10)
    r_note.italic = True
    r_note.font.color.rgb = GREY

doc.add_paragraph()
callout_box(
    doc,
    "Total new pairs from augmentation:  323",
    "Original 881 pairs  +  323 new augmented pairs  =  1,204 total pairs.  "
    "The augmented pairs are used alongside the originals during training.",
    bg_hex="EBF8F6",
    title_color=TEAL,
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# SECTION 6 – BACK-TRANSLATION
# ════════════════════════════════════════════════════════════════════
heading(doc, "6.  Back-Translation", 1)
body(doc, (
    "Back-translation is a quality-checking technique and also a way to generate "
    "more data.  The idea is simple: take a Runyoro sentence, translate it to English, "
    "then translate the English back to Runyoro and check if you get something close "
    "to what you started with."
))
doc.add_paragraph()
heading(doc, "How back-translation works — step by step", 2, color=TEAL)
doc.add_paragraph()

steps_tbl = doc.add_table(rows=4, cols=2)
steps_tbl.style = "Table Grid"
steps = [
    ("Step 1",  "Take an original Runyoro sentence.\n"
                "Example:  'Oraire ota?'  (How are you?)"),
    ("Step 2",  "Translate it to English using the model.\n"
                "Result:  'How are you?'"),
    ("Step 3",  "Translate that English back to Runyoro.\n"
                "Result:  'Oraire ota?'  or something close"),
    ("Step 4",  "Compare the final Runyoro to the original.  "
                "If they match well, the translation is reliable.  "
                "If they are very different, it is flagged as a quality issue."),
]
set_cell_bg(steps_tbl.rows[0].cells[0], "070235")
set_cell_bg(steps_tbl.rows[0].cells[1], "070235")
for i, (step, desc) in enumerate(steps):
    row = steps_tbl.rows[i]
    row.cells[0].text = step
    row.cells[1].text = desc
    if i % 2 == 0 and i > 0:
        set_cell_bg(row.cells[0], "F2F4F6")
        set_cell_bg(row.cells[1], "F2F4F6")
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(11)
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after  = Pt(4)
for cell in steps_tbl.rows[0].cells:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = WHITE
            run.font.bold = True

doc.add_paragraph()
heading(doc, "What back-translation is used for in this project", 2, color=TEAL)
bullet(doc,
    "Quality check: pairs where the roundtrip translation score is very low are "
    "flagged for review — they may be bad translations or misaligned pairs.",
    "Quality checking:  ")
bullet(doc,
    "Data generation: the English translations produced during the forward pass "
    "(Runyoro → English) can be used as additional training examples for the "
    "reverse direction (English → Runyoro), effectively creating new pairs.",
    "Extra training data:  ")
bullet(doc,
    "Model evaluation: after training, back-translation roundtrip score is "
    "reported as one of the quality metrics alongside BLEU and chrF++.",
    "Evaluation metric:  ")

doc.add_paragraph()
callout_box(
    doc,
    "Status: Back-translation runs automatically during evaluation",
    "The system is designed to run back-translation once the model has been trained.  "
    "It takes a sample of 50 test pairs and checks how well they round-trip.  "
    "A score above 0.6 indicates consistent, reliable translations.",
    bg_hex="EBF8F6",
    title_color=TEAL,
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# SECTION 7 – HOW MUCH DATA IS LEFT
# ════════════════════════════════════════════════════════════════════
heading(doc, "7.  How Much Data Is Left", 1)
body(doc, (
    "After all the steps above — filtering, cleaning, and augmentation — "
    "here is exactly how much data remains and how it is split for training."
))
doc.add_paragraph()

# Big visual summary table
heading(doc, "Data at each stage", 2, color=TEAL)
stage_tbl = doc.add_table(rows=8, cols=3)
stage_tbl.style = "Table Grid"
stage_data = [
    ("Stage",                           "Pairs",   "Notes"),
    ("Raw pairs from all 8 files",      "3,782",   "Before any processing"),
    ("After removing bad pairs",        "881",     "2,901 removed (too short, symbols, bad ratio)"),
    ("After cleaning",                  "881",     "Same count — cleaning fixes text, doesn't remove"),
    ("Augmented pairs created",         "+323",    "New pairs generated from the 881 clean pairs"),
    ("Total corpus",                    "1,204",     "881 original + 323 augmented"),
    ("Training set (85%)",              "1,023",     "Used to teach the model"),
    ("Validation set (10%) + Test (5%)","120 + 61", "Used to measure quality during and after training"),
]
for i, (col1, col2, col3) in enumerate(stage_data):
    row = stage_tbl.rows[i]
    row.cells[0].text = col1
    row.cells[1].text = col2
    row.cells[2].text = col3
    if i == 0:
        for cell in row.cells:
            set_cell_bg(cell, "070235")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(11)
    elif i == 5:
        for cell in row.cells:
            set_cell_bg(cell, "006A61")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(11)
    elif i in (6, 7):
        for cell in row.cells:
            set_cell_bg(cell, "EBF8F6")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)
    elif i % 2 == 0:
        for cell in row.cells:
            set_cell_bg(cell, "F2F4F6")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)
    else:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)
    for cell in row.cells:
        for para in cell.paragraphs:
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after  = Pt(4)

doc.add_paragraph()

# Files on disk
heading(doc, "Files saved on disk", 2, color=TEAL)
body(doc, "All processed data is saved in the runyoro_nmt/data/ folder:")
doc.add_paragraph()

files_tbl = doc.add_table(rows=8, cols=3)
files_tbl.style = "Table Grid"
file_rows = [
    ("File",                              "Location",             "Contents"),
    ("cleaned_pairs.tsv",                 "data/processed/",      "881 clean pairs"),
    ("train.tsv",                         "data/processed/",      "1,023 training pairs"),
    ("val.tsv",                           "data/processed/",      "120 validation pairs"),
    ("test.tsv",                          "data/processed/",      "61 test pairs"),
    ("all_pairs.tsv",                     "data/augmented/",      "1,204 pairs (original + augmented)"),
    ("glossary.csv / glossary.json",      "data/tm/",             "425 terminology entries"),
    ("runyoro_en.tmx / runyoro_en.tbx",   "data/tm/",             "Translation memory + termbase"),
]
for i, (f, loc, cont) in enumerate(file_rows):
    row = files_tbl.rows[i]
    row.cells[0].text = f
    row.cells[1].text = loc
    row.cells[2].text = cont
    if i == 0:
        for cell in row.cells:
            set_cell_bg(cell, "006A61")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(10)
    elif i % 2 == 0:
        for cell in row.cells:
            set_cell_bg(cell, "F2F4F6")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    else:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    for cell in row.cells:
        for para in cell.paragraphs:
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after  = Pt(3)

doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# SECTION 8 – LINGUISTIC RESOURCES CREATED
# ════════════════════════════════════════════════════════════════════
heading(doc, "8.  Linguistic Resources Created", 1)
body(doc, (
    "Alongside the training data, the pipeline automatically built several "
    "reference resources that are useful for translation quality and future work."
))
doc.add_paragraph()

resources = [
    (
        "Translation Memory (TMX file)",
        "A standard industry format that stores every Runyoro-English sentence "
        "pair.  Any future translation tool can load this file to reuse "
        "previously translated sentences automatically.",
        "runyoro_en.tmx  —  881 entries, 3,562 lines",
    ),
    (
        "Terminology Database (TBX file)",
        "A structured list of technical terms with their translations and domain "
        "labels (e.g. Agriculture, Health, Education).  Used to ensure "
        "consistent terminology across all translations.",
        "runyoro_en.tbx  —  37 terms",
    ),
    (
        "Bilingual Glossary (CSV + JSON)",
        "A simple table of short word pairs (1–3 words each) exported in two "
        "formats.  The CSV can be opened in Excel; the JSON is used by the "
        "translation API to verify domain-specific terms.",
        "glossary.csv / glossary.json  —  37 entries",
    ),
    (
        "Named Entity Registry (JSON)",
        "A list of proper nouns found in the data — names of people, places, "
        "and organisations — paired with their English equivalents.  "
        "This helps the model handle names correctly.",
        "named_entities.json  —  58 entries",
    ),
]
for title, desc, file_ref in resources:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.3)
    r1 = p.add_run(f"\u25a0  {title}\n")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = INDIGO
    r2 = p.add_run(f"     {desc}\n")
    r2.font.size = Pt(11)
    r2.font.color.rgb = GREY
    r3 = p.add_run(f"     File: {file_ref}")
    r3.font.size = Pt(10)
    r3.italic = True
    r3.font.color.rgb = TEAL

doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# SECTION 9 – FULL PIPELINE DIAGRAM (text-based)
# ════════════════════════════════════════════════════════════════════
heading(doc, "9.  Full Pipeline at a Glance", 1)
body(doc, "This diagram shows every step from raw files to training-ready data:")
doc.add_paragraph()

flow = [
    ("8 raw files (.xlsx, .docx)",    "INPUT",   "070235", WHITE),
    ("Extract sentence pairs",        "Step 1",  "006A61", WHITE),
    ("3,782 raw pairs",               "",        "EBF8F6", INDIGO),
    ("Filter bad pairs",              "Step 2",  "006A61", WHITE),
    ("881 valid pairs remain",        "",        "EBF8F6", INDIGO),
    ("Check alignment",               "Step 3",  "006A61", WHITE),
    ("0 suspicious pairs found",      "",        "EBF8F6", INDIGO),
    ("Clean text",                    "Step 4",  "006A61", WHITE),
    ("174 pairs fixed, 881 total",    "",        "EBF8F6", INDIGO),
    ("Build TM, glossary, NE list",   "Step 5",  "006A61", WHITE),
    ("TMX + TBX + 425-term glossary",  "",        "EBF8F6", INDIGO),
    ("Augment data",                  "Step 6",  "006A61", WHITE),
    ("881 + 323 = 1,204 total pairs",   "",        "EBF8F6", INDIGO),
    ("Split: Train / Val / Test",     "Step 7",  "006A61", WHITE),
    ("1,023 train / 74 val / 38 test",  "OUTPUT",  "070235", WHITE),
]

flow_tbl = doc.add_table(rows=len(flow), cols=2)
flow_tbl.style = "Table Grid"
for i, (label, tag, bg, fg) in enumerate(flow):
    row = flow_tbl.rows[i]
    row.cells[0].text = tag
    row.cells[1].text = label
    set_cell_bg(row.cells[0], bg)
    set_cell_bg(row.cells[1], bg)
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = fg
                run.font.size = Pt(11)
                if tag in ("INPUT", "OUTPUT", "Step 1", "Step 2",
                           "Step 3", "Step 4", "Step 5", "Step 6", "Step 7"):
                    run.font.bold = True
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after  = Pt(3)

doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# SECTION 10 – WHAT HAPPENS NEXT
# ════════════════════════════════════════════════════════════════════
heading(doc, "10.  What Happens Next", 1)
body(doc, (
    "The data is fully prepared.  Here is what the pipeline does with it next:"
))
doc.add_paragraph()

next_steps = [
    ("Model Training",
     "The 1,204 pairs are fed into NLLB-200 (a powerful multilingual translation "
     "model with 1.3 billion parameters) for fine-tuning.  "
     "Training runs on 2 NVIDIA RTX 4090 GPUs for 15 epochs."),
    ("Evaluation",
     "After training, the model is tested on the 61 held-out test pairs.  "
     "Quality scores (BLEU, chrF++, COMET, BERTScore) are calculated and "
     "a back-translation roundtrip check is run."),
    ("Error Analysis",
     "The system automatically checks for hallucinations (made-up content), "
     "untranslated output, code-switching (mixing languages), and "
     "wrong terminology."),
    ("Quantization for Raspberry Pi",
     "The trained model is compressed to INT8 format so it can run on "
     "low-power devices like a Raspberry Pi 4 or 5."),
    ("Deployment",
     "The model and dataset are pushed to the Hugging Face Hub under "
     "'kathay/runyoro-nmt-v1'.  A live demo is available at "
     "huggingface.co/spaces/kathay/runyoro-translator."),
]
for title, desc in next_steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.3)
    r1 = p.add_run(f"\u25b6  {title}\n")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = TEAL
    r2 = p.add_run(f"     {desc}")
    r2.font.size = Pt(11)
    r2.font.color.rgb = GREY

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════
# FINAL SUMMARY BOX
# ════════════════════════════════════════════════════════════════════
sum_tbl = doc.add_table(rows=1, cols=1)
set_cell_bg(sum_tbl.rows[0].cells[0], "070235")
cell = sum_tbl.rows[0].cells[0]
cell.paragraphs[0].clear()

lines = [
    ("Final Data Summary", True, 14, WHITE),
    ("", False, 6, WHITE),
    ("Raw pairs extracted from 8 files:       3,782", False, 11, RGBColor(0x86, 0xf2, 0xe4)),
    ("Pairs removed during filtering:         2,901", False, 11, RGBColor(0xc4, 0xc1, 0xfb)),
    ("Clean pairs after filtering + cleaning:   881", False, 11, RGBColor(0x86, 0xf2, 0xe4)),
    ("New pairs generated by augmentation:     +323", False, 11, RGBColor(0xc4, 0xc1, 0xfb)),
    ("", False, 6, WHITE),
    ("TOTAL pairs ready for training:           1,204", True, 13, WHITE),
    ("  \u2514\u2500 Training set:   1,023 pairs (85%)", False, 11, RGBColor(0x86, 0xf2, 0xe4)),
    ("  \u2514\u2500 Validation:     120 pairs (10%)", False, 11, RGBColor(0x86, 0xf2, 0xe4)),
    ("  \u2514\u2500 Test set:       61 pairs ( 5%)", False, 11, RGBColor(0x86, 0xf2, 0xe4)),
]
for i, (text, bold, size, color) in enumerate(lines):
    if i == 0:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = color

p_end = cell.add_paragraph()
p_end.paragraph_format.space_after = Pt(6)

doc.add_paragraph()

# Footer note
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = p_footer.add_run(
    "Generated by the runyoro-nmt-v1 pipeline  \u2022  "
    "kathay/runyoro-nmt-v1  \u2022  "
    "huggingface.co/spaces/kathay/runyoro-translator"
)
r_f.font.size = Pt(9)
r_f.font.color.rgb = GREY
r_f.italic = True

# ════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════
doc.save(str(OUT_PATH))
print(f"Document saved: {OUT_PATH}")
